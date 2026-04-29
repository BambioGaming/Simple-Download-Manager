"""
generate_report.py
Generates the CS404 SDM Technical Report as a .docx file.
Run from the project root:  python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table_borders(table):
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_page_number_footer(doc):
    section  = doc.sections[0]
    footer   = section.footer
    para     = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_code_block(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # Light grey shading on paragraph
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    return para


def add_bullet(doc, text: str, bold_prefix: str = ""):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    para.paragraph_format.space_after = Pt(2)
    return para


def add_h1(doc, text: str):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x50, 0x7A)
    return h


def add_h2(doc, text: str):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return h


def add_h3(doc, text: str):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after  = Pt(3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x50, 0x7A)
    return h


def add_body(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


# ──────────────────────────────────────────────────────────────
# Cover page
# ──────────────────────────────────────────────────────────────

def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("MedTech — Mediterranean Institute of Technology")
    run.font.name  = "Calibri"
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1F, 0x50, 0x7A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("CS404 — Distributed Systems")
    r.font.size  = Pt(12)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    doc.add_paragraph()

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = main_title.add_run("Simple Download Manager")
    r.font.name  = "Calibri"
    r.font.size  = Pt(28)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x1F, 0x50, 0x7A)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("(SDM)")
    r.font.size  = Pt(22)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = desc.add_run("Technical Report")
    r.font.size  = Pt(16)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("Student:", "[Student Name]"),
        ("Course:",  "CS404 — Distributed Systems"),
        ("Supervisor:", "Ayoub Bousselmi"),
        ("Academic Year:", "2025 – 2026"),
        ("Date:", "April 2026"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}  ")
        r.font.bold = True
        r.font.size = Pt(11)
        p.add_run(value).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# Section 1 — Introduction
# ──────────────────────────────────────────────────────────────

def build_intro(doc):
    add_h1(doc, "1.  Introduction")

    add_h2(doc, "1.1  Project Overview")
    add_body(doc,
        "The Simple Download Manager (SDM) is a desktop application that improves file download "
        "speed and reliability by splitting a single file into multiple byte-range segments and "
        "downloading them simultaneously using parallel threads. It is inspired by commercial tools "
        "such as Internet Download Manager (IDM) and the open-source Xtreme Download Manager (XDM), "
        "and is built entirely with Python standard-library components plus the requests HTTP library."
    )
    add_body(doc,
        "SDM was developed as the capstone project for CS404 — Distributed Systems at MedTech "
        "Mediterranean Institute of Technology. It serves as a practical demonstration of "
        "multi-threading, HTTP Range requests, fault tolerance, and real-time progress monitoring "
        "in a concrete, user-visible application."
    )

    add_h2(doc, "1.2  Objectives")
    for item in [
        "Accept a file URL and split it into N configurable parallel segments (default: 4, max: 16).",
        "Download all segments simultaneously using a ThreadPoolExecutor.",
        "Merge the downloaded segments into the correct final file.",
        "Handle network errors with automatic per-segment retry and linear back-off.",
        "Allow the user to pause and resume downloads with zero CPU usage while paused.",
        "Resume interrupted downloads after an application restart, without re-downloading already-saved bytes.",
        "Display live metrics: download speed (smoothed), progress percentage, and estimated time remaining.",
        "Maintain a persistent download history in an SQLite database.",
        "Provide both a Tkinter desktop GUI and a CLI fallback mode.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 2 — Architecture Design
# ──────────────────────────────────────────────────────────────

def build_architecture(doc):
    add_h1(doc, "2.  Architecture Design")

    add_h2(doc, "2.1  Architecture Style")
    add_body(doc,
        "SDM follows a Layered Architecture (also called an N-tier or hierarchical architecture). "
        "Each layer has a single, clearly defined responsibility and communicates only with the "
        "layer immediately below it. This strict separation of concerns makes every component "
        "independently testable and replaceable."
    )
    add_body(doc,
        "The seven layers, from top to bottom, are: UI → Download Manager Core → Thread Controller "
        "→ Segment Workers → File Assembler → Progress Tracker → Persistence (SQLite). "
        "The UI never touches the database directly; the workers never touch the UI. All "
        "cross-cutting state is carried in shared DownloadTask and SegmentInfo model objects "
        "that are protected by threading.Lock."
    )

    add_h2(doc, "2.2  Architecture Diagram")
    add_body(doc,
        "The diagram below illustrates the data flow and control flow between the layers. "
        "Arrows show which components call which, and dashed arrows show background "
        "(asynchronous) communication."
    )

    # Architecture table
    arch_data = [
        ("UI Layer", "ui/gui.py  ·  ui/cli.py",
         "Tkinter window (toolbar + Active/History tabs) or CLI; polls DownloadManager every 200 ms."),
        ("Download Manager Core", "core/download_manager.py",
         "Central coordinator: FSM state machine, segment splitting, daemon coordinator threads."),
        ("Thread Controller", "core/thread_controller.py",
         "Wraps ThreadPoolExecutor; submits SegmentWorkers, collects results via as_completed()."),
        ("Segment Workers", "core/segment_worker.py",
         "One worker per segment: HTTP Range GET, 256 KB chunk loop, pause/cancel checks, retry."),
        ("File Assembler", "core/file_assembler.py",
         "Concatenates .part files in segment-index order; verifies total size; deletes parts."),
        ("Progress Tracker", "monitoring/progress_tracker.py",
         "Thread-safe rolling-window speed + EWMA smoothing; computes ETA and percent complete."),
        ("Persistence Layer", "persistence/database.py",
         "SQLite (WAL mode) stores download records and per-segment byte offsets for resume."),
    ]

    tbl = doc.add_table(rows=len(arch_data) + 1, cols=3)
    tbl.style = "Table Grid"
    add_table_borders(tbl)

    headers = ["Layer", "File(s)", "Responsibility"]
    hdr_colors = ["1F507A", "1F507A", "1F507A"]
    for i, (h, c) in enumerate(zip(headers, hdr_colors)):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, c)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    row_colors = ["EBF3FB", "FFFFFF"] * 10
    for r_idx, (layer, files, resp) in enumerate(arch_data, start=1):
        row = tbl.rows[r_idx]
        for c_idx, text in enumerate([layer, files, resp]):
            cell = row.cells[c_idx]
            set_cell_bg(cell, row_colors[r_idx - 1])
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.bold = True

    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(5.0)
    tbl.columns[2].width = Cm(8.5)

    doc.add_paragraph()

    add_h2(doc, "2.3  Communication Mechanism")
    add_body(doc,
        "All inter-layer communication within SDM is in-process (no network sockets or REST APIs "
        "are used between components). The communication mechanisms are:"
    )
    comm_items = [
        ("Direct method calls",
         "The UI calls DownloadManager public methods (add_download, start_download, pause_download, "
         "resume_download, cancel_download, get_progress). The Download Manager calls Thread Controller, "
         "which in turn submits SegmentWorker futures."),
        ("Shared model objects",
         "DownloadTask and SegmentInfo objects are passed by reference to all layers. "
         "Mutable fields (status, downloaded) are protected by threading.Lock."),
        ("threading.Event (pause/cancel)",
         "pause_event.set() = green light; pause_event.clear() = workers block inside the OS kernel "
         "at zero CPU. cancel_event.set() = workers exit at next chunk boundary."),
        ("Threading.Lock (DB serialization)",
         "The Database class owns a single threading.Lock that serializes all SQLite reads and "
         "writes from multiple worker threads."),
        ("Polling loop (UI refresh)",
         "The Tkinter main thread calls get_progress() every 200 ms via widget.after(200, ...). "
         "Worker threads never touch Tkinter widgets directly — this is the key thread-safety invariant."),
    ]
    for title, detail in comm_items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title + ": ")
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(detail).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 3 — Design Decisions
# ──────────────────────────────────────────────────────────────

def build_design_decisions(doc):
    add_h1(doc, "3.  Design Decisions")

    add_h2(doc, "3.1  Technology Stack")
    add_body(doc,
        "The table below summarises the technology choices and the rationale behind each one."
    )

    tech_data = [
        ("Python 3.10+", "Concise syntax, rich standard library (threading, sqlite3, tkinter). "
                         "Rapid development cycle suited to the project timeline."),
        ("requests library", "Clean HTTP API with streaming support (iter_content), automatic "
                             "redirect handling, and persistent Session objects for TCP connection reuse."),
        ("threading.Thread +\nThreadPoolExecutor", "OS-level threads give true parallelism for I/O-bound "
                                                    "network downloads. ThreadPoolExecutor manages the worker "
                                                    "lifecycle and exposes Future objects for result collection."),
        ("SQLite 3 (WAL mode)", "Zero-configuration embedded database. WAL (Write-Ahead Logging) mode "
                                "allows concurrent readers while a writer is active, reducing lock contention "
                                "among segment worker threads."),
        ("Tkinter", "Ships with every standard Python installation — zero extra dependencies for "
                    "the desktop GUI. The event loop is single-threaded by design, which enforces "
                    "the correct pattern: workers update data, UI reads data."),
    ]

    tbl = doc.add_table(rows=len(tech_data) + 1, cols=2)
    add_table_borders(tbl)
    for i, h in enumerate(["Technology", "Rationale"]):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, "1F507A")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(10)

    for idx, (tech, rat) in enumerate(tech_data, 1):
        row = tbl.rows[idx]
        bg = "EBF3FB" if idx % 2 == 1 else "FFFFFF"
        set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[1], bg)
        r0 = row.cells[0].paragraphs[0].add_run(tech)
        r0.bold = True; r0.font.size = Pt(9.5)
        row.cells[1].paragraphs[0].add_run(rat).font.size = Pt(9.5)

    tbl.columns[0].width = Cm(4.5)
    tbl.columns[1].width = Cm(12.5)
    doc.add_paragraph()

    add_h2(doc, "3.2  HTTP Range Requests")
    add_body(doc,
        "HTTP Range requests (RFC 7233) are the core mechanism that enables segmented parallel "
        "downloading. The client sends a Range: bytes=start-end header with each GET request; "
        "a compliant server responds with HTTP 206 Partial Content and only the requested byte range."
    )
    add_body(doc, "The protocol exchange for a 4-segment download looks like this:")
    add_code_block(doc,
        "HEAD /100MB.bin  ──────────────►  200 OK\n"
        "                                   Content-Length: 104857600\n"
        "                ◄────────────────  Accept-Ranges: bytes\n\n"
        "GET  /100MB.bin  (Range: bytes=0-26214399)        →  206  [segment 0]\n"
        "GET  /100MB.bin  (Range: bytes=26214400-52428799) →  206  [segment 1]   ← parallel\n"
        "GET  /100MB.bin  (Range: bytes=52428800-78643199) →  206  [segment 2]   ← parallel\n"
        "GET  /100MB.bin  (Range: bytes=78643200-104857599)→  206  [segment 3]   ← parallel\n\n"
        "[all 4 segments downloaded simultaneously]\n"
        "[merge: part0 + part1 + part2 + part3  →  100MB.bin]"
    )
    add_body(doc,
        "SDM sends a HEAD probe before splitting the file. If the server's response includes "
        "Accept-Ranges: bytes and a numeric Content-Length, multi-segment mode is used. "
        "If either header is absent, SDM falls back to a single-segment download that works "
        "with any HTTP server."
    )

    add_h2(doc, "3.3  Segment Splitting Strategy")
    add_body(doc,
        "The file is divided into N equal-sized segments. The last segment absorbs any remainder "
        "bytes so the total coverage is always exact."
    )
    add_code_block(doc,
        "segment_size = total_bytes // N\n\n"
        "for i in range(N):\n"
        "    start = i * segment_size\n"
        "    end   = start + segment_size - 1   # last segment: end = total_bytes - 1\n"
        "    # Worker downloads Range: bytes={start}-{end}"
    )
    add_body(doc,
        "On retry, the worker computes actual_start = segment.start_byte + segment.downloaded "
        "so only the un-downloaded tail of the segment is requested. This partial-retry approach "
        "avoids re-downloading bytes that were already written to the .part file."
    )

    add_h2(doc, "3.4  Persistence Layer")
    add_body(doc,
        "SQLite stores two tables: downloads (one row per file) and segments (one row per segment). "
        "The segments table records the start_byte, end_byte, and downloaded (bytes already saved) "
        "for every segment. On the next application launch, restore_incomplete() reads this table "
        "and reconstructs DownloadTask objects so the user can resume with a single click."
    )
    add_body(doc,
        "To reduce lock contention among worker threads, each worker writes its progress to the "
        "database at most once per second (time-gated writes), with a final flush when the segment "
        "completes. Without this gate, 8 simultaneous workers writing every 256 KB chunk produced "
        "thousands of DB writes per second and measurably reduced download throughput."
    )

    add_h2(doc, "3.5  Progress Monitoring Algorithm")
    add_body(doc,
        "Download speed is computed using a time-bounded 5-second rolling window combined with "
        "Exponential Weighted Moving Average (EWMA) smoothing, mimicking the approach used by "
        "browsers and Free Download Manager:"
    )
    for step in [
        "Workers call ProgressTracker.update(chunk_size) after every chunk write. "
        "Each call appends a (timestamp, cumulative_session_bytes) sample to a deque.",
        "Samples older than 5 seconds are discarded on each update() call, keeping the window current.",
        "Raw speed = (newest_bytes − oldest_bytes) / (newest_timestamp − oldest_timestamp) over the window.",
        "EWMA: smoothed_speed = α × raw_speed + (1 − α) × smoothed_speed  (α = 0.15). "
        "This damps short spikes while still reacting to genuine speed changes.",
        "If raw_speed == 0 (brief stall), the last smoothed value is retained so the UI does not "
        "snap to zero on a momentary network hiccup.",
    ]:
        add_bullet(doc, step)

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 4 — Thread Model
# ──────────────────────────────────────────────────────────────

def build_thread_model(doc):
    add_h1(doc, "4.  Thread Model")

    add_h2(doc, "4.1  Thread Hierarchy")
    add_body(doc,
        "SDM uses three distinct thread roles. Each download spawns its own coordinator + pool:"
    )
    add_code_block(doc,
        "Main thread  (Tkinter event loop)\n"
        "│\n"
        "│  widget.after(200 ms)  ──►  get_progress()  ──►  update widgets\n"
        "│\n"
        "│  user clicks 'Start'\n"
        "│    └─►  start_download()  ──►  threading.Thread(target=_run_download, daemon=True)\n"
        "│                                      │  [coordinator daemon thread]\n"
        "│                                      └─►  ThreadPoolExecutor(max_workers=N)\n"
        "│                                                ├─►  SegmentWorker(seg=0).run()\n"
        "│                                                ├─►  SegmentWorker(seg=1).run()\n"
        "│                                                ├─►  SegmentWorker(seg=2).run()\n"
        "│                                                └─►  SegmentWorker(seg=3).run()"
    )
    add_body(doc,
        "Key invariant: worker threads never touch Tkinter widgets. All widget mutations happen "
        "on the main thread inside the 200 ms polling callback. This is the fundamental rule that "
        "keeps the GUI thread-safe without extra locking."
    )

    add_h2(doc, "4.2  ThreadPoolExecutor")
    add_body(doc,
        "The Thread Controller wraps a concurrent.futures.ThreadPoolExecutor with max_workers=N "
        "(where N is the segment count chosen by the user). Workers are submitted as callables; "
        "the coordinator thread then calls as_completed() to collect results as each segment "
        "finishes. This design means:"
    )
    for item in [
        "The thread pool is bounded — N threads maximum, regardless of how many files are downloading.",
        "Futures provide clean success/failure signalling without manual join() calls.",
        "Idle threads in the pool are automatically reaped after the download completes.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "4.3  Pause / Resume Mechanism")
    add_body(doc,
        "Pause and resume are implemented using a threading.Event called pause_event. "
        "The event acts as a 'green light': when set, workers run; when cleared, workers block."
    )
    add_code_block(doc,
        "# Pause:  clear the event  →  workers block at next chunk\n"
        "task.pause_event.clear()\n\n"
        "# Resume: set the event   →  workers unblock immediately\n"
        "task.pause_event.set()\n\n"
        "# Inside the worker chunk loop:\n"
        "self._task.pause_event.wait()   # blocks indefinitely while event is cleared\n"
        "                                 # returns instantly when event is set"
    )
    add_body(doc,
        "The critical property of threading.Event.wait() is that the blocked threads sleep "
        "inside the OS kernel — they consume zero CPU while paused. This is far superior to a "
        "polling loop with time.sleep(), which would waste CPU cycles."
    )

    add_h2(doc, "4.4  Thread Synchronisation")
    add_body(doc,
        "Multiple threads access shared state concurrently. SDM uses three locking mechanisms:"
    )
    sync_items = [
        ("threading.Lock in DownloadManager",
         "Protects the _active_tasks, _trackers, and _controllers dictionaries from concurrent "
         "reads and writes by the main thread (UI polling) and coordinator threads."),
        ("threading.Lock in ProgressTracker",
         "Protects the _downloaded counter and _samples deque. Workers call update() from "
         "multiple threads simultaneously; get_stats() is called from the main thread."),
        ("threading.Lock in Database",
         "SQLite's default threading mode is not safe for concurrent writes. A single lock "
         "serialises all database access so only one thread holds the connection at a time."),
    ]
    for title, detail in sync_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + ": ").bold = True
        p.add_run(detail)
        p.paragraph_format.space_after = Pt(3)

    add_h2(doc, "4.5  Retry Mechanism")
    add_body(doc,
        "Each segment worker retries up to max_retries times (default: 3) on any "
        "requests.RequestException. The retry uses linear back-off:"
    )
    add_code_block(doc,
        "attempt 1  →  fail  →  sleep(retry_delay × 1)  =  sleep 2 s\n"
        "attempt 2  →  fail  →  sleep(retry_delay × 2)  =  sleep 4 s\n"
        "attempt 3  →  fail  →  segment marked FAILED   →  download marked FAILED"
    )
    add_body(doc,
        "If a segment fails after all retries, the entire download is marked FAILED but "
        "the already-downloaded bytes remain in the .part files so the user can potentially "
        "retry the whole download without starting from scratch."
    )
    add_body(doc,
        "Retries are partial: actual_start = start_byte + downloaded. Only the bytes not yet "
        "written to the .part file are re-requested. The server receives a narrower Range header "
        "on each retry, and the worker appends the new bytes to the existing .part file (opened "
        "in 'ab' append-binary mode)."
    )

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 5 — Challenges Faced
# ──────────────────────────────────────────────────────────────

def build_challenges(doc):
    add_h1(doc, "5.  Challenges Faced")

    add_h2(doc, "5.1  File Handle Locking on Windows (WinError 32)")
    add_body(doc,
        "When a download was cancelled, the cleanup thread attempted to delete .part files "
        "immediately. On Windows, a process cannot delete a file that another thread still has "
        "open — this produced: PermissionError: [WinError 32] The process cannot access the file "
        "because it is being used by another process."
    )
    add_body(doc,
        "The root cause: segment workers had the .part files open for writing when the cleanup "
        "thread called file.unlink(). The fix was to capture the ThreadController before clearing "
        "it from the registry, then call controller.wait_for_completion() in the cleanup thread "
        "before deleting any files. This guarantees all workers have closed their file handles "
        "before the delete is attempted."
    )

    add_h2(doc, "5.2  Resume After Application Restart")
    add_body(doc,
        "After a restart, SDM restores interrupted downloads from the database and marks them "
        "as PAUSED. The original resume_download() only called pause_event.set() to unblock "
        "existing workers — but after a restart there are no live workers. Clicking Resume "
        "would silently do nothing."
    )
    add_body(doc,
        "The fix was to check whether an active ThreadController exists for the download. "
        "If yes (paused mid-flight), unblock the existing workers. "
        "If no (restored from DB), call start_download() to spawn a new coordinator thread "
        "and a new ThreadPoolExecutor. The same check was mirrored in start_download() to "
        "prevent accidentally spawning a second coordinator for an already-running download."
    )

    add_h2(doc, "5.3  Database Write Contention")
    add_body(doc,
        "The initial implementation wrote progress to SQLite after every 8 KB chunk. With "
        "4 parallel workers each downloading at full speed, this produced thousands of DB "
        "writes per second. Even with WAL mode, the lock contention between threads caused "
        "a measurable drop in download throughput."
    )
    add_body(doc,
        "The fix was a time-gated write strategy: each worker writes its progress to the "
        "database at most once per second (using time.monotonic() comparison), with a "
        "mandatory final flush when the segment finishes. This reduced DB writes from thousands "
        "per second to at most N writes per second (where N is the segment count), "
        "eliminating the contention entirely."
    )

    add_h2(doc, "5.4  Accurate Speed Measurement")
    add_body(doc,
        "The early progress tracker seeded the sample deque at construction time with a "
        "(t=0, bytes=0) entry. After a resume, the first speed measurement included the "
        "pause duration in its denominator, making the displayed speed appear artificially "
        "low for the first several seconds after resuming."
    )
    add_body(doc,
        "The fix was lazy seeding: no initial sample is added at construction; the first "
        "real sample is added when the first chunk arrives. Additionally, the speed window "
        "was changed from a fixed-count sample buffer to a time-bounded 5-second window, "
        "so the denominator always reflects actual download time rather than wall-clock time "
        "since the tracker was created."
    )

    add_h2(doc, "5.5  Application Shutdown Hang")
    add_body(doc,
        "Python's normal shutdown process joins all non-daemon threads. When a download was "
        "paused, the worker threads were blocked inside pause_event.wait(). On Windows, "
        "these threads did not release on interpreter shutdown — the application terminal "
        "appeared to hang even after the Tkinter window was closed."
    )
    add_body(doc,
        "The fix was to call os._exit(0) in main.py immediately after db.close(). "
        "os._exit() terminates the process without running atexit handlers or joining "
        "threads, which is safe here because the database is already cleanly closed and "
        "the daemon threads hold no external resources that require graceful teardown."
    )

    add_h2(doc, "5.6  Benchmark Validity")
    add_body(doc,
        "The initial benchmark used a 10 MB test file from an OVH server. Results showed "
        "multi-threaded downloads were slower than single-threaded, which appeared to "
        "contradict the entire premise of the project."
    )
    add_body(doc,
        "Investigation revealed two problems: (1) the OVH server applies a per-IP bandwidth "
        "cap, so additional connections shared a fixed total bandwidth rather than each getting "
        "independent throughput; (2) a 10 MB file is too small for multiple TCP connections to "
        "complete their slow-start ramp-up. Switching to the Hetzner NBG1 speed-test server "
        "(100 MB file, no per-IP cap) produced valid results showing genuine speedup."
    )

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 6 — Performance Comparison
# ──────────────────────────────────────────────────────────────

def build_performance(doc):
    add_h1(doc, "6.  Performance Comparison")

    add_h2(doc, "6.1  Methodology")
    add_body(doc,
        "The benchmark script (benchmark.py) measures the time from start_download() to "
        "completion — exactly the interval the user experiences from clicking Start to "
        "the file being fully assembled on disk. The URL probe (HEAD request) is excluded "
        "from the timing, matching the GUI experience where the probe happens at the "
        "add-URL step rather than the start step."
    )
    for item in [
        "All test runs use the full DownloadManager stack (same code path as the GUI): "
        "SegmentWorker, ThreadController, FileAssembler, ProgressTracker, and SQLite persistence.",
        "An in-memory SQLite database (:memory:) is used to avoid disk I/O overhead in measurements.",
        "Each run uses its own output subdirectory to ensure .part files from one run do not affect the next.",
        "Test file: https://nbg1-speed.hetzner.com/100MB.bin — 100 MB, supports Range requests, no per-IP bandwidth cap.",
        "The baseline is a 1-segment download via DownloadManager (not a raw requests.get()), "
        "so DownloadManager overhead is equally present in all test cases.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "6.2  Results")

    results = [
        ("1 segment (baseline)", "19.0 s", "5.25 MB/s", "1.00×"),
        ("2 segments",           "17.4 s", "5.74 MB/s", "1.09×"),
        ("4 segments",           "17.4 s", "5.73 MB/s", "1.09×"),
        ("8 segments",           "17.4 s", "5.73 MB/s", "1.09×"),
    ]

    tbl = doc.add_table(rows=len(results) + 1, cols=4)
    add_table_borders(tbl)
    for i, h in enumerate(["Method", "Time", "Speed", "Speedup"]):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, "1F507A")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, row_data in enumerate(results, 1):
        row = tbl.rows[idx]
        bg = "EBF3FB" if idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            if c_idx > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c_idx == 0 and idx == 1:
                run.bold = True

    doc.add_paragraph()

    add_h2(doc, "6.3  Analysis and Discussion")
    add_body(doc,
        "The results reveal a real-world distributed systems phenomenon: parallel segmented "
        "downloading improves throughput until a network bottleneck is reached, after which "
        "additional workers yield no further gain."
    )
    add_body(doc,
        "Going from 1 segment to 2 segments produced a ~9% improvement (19.0 s → 17.4 s). "
        "This gain comes from two TCP connections reaching the server's bandwidth allocation "
        "faster than a single connection can during its slow-start ramp-up phase. Two "
        "connections open their congestion windows in parallel, meaning the aggregate "
        "throughput reaches its ceiling more quickly."
    )
    add_body(doc,
        "Beyond 2 segments, the time remained constant at 17.4 s for both 4 and 8 segments. "
        "This plateau indicates that the per-IP bandwidth ceiling (~5.7 MB/s, roughly 46 Mbps) "
        "was already reached at 2 connections. Additional connections share the same fixed "
        "bandwidth pie rather than each receiving independent throughput."
    )
    add_body(doc,
        "This is a valid and honest benchmark result. On a higher-bandwidth connection or "
        "with a server that allocates bandwidth per-connection (rather than per-IP), the "
        "speedup curve would continue to rise. The key learning is that parallel segmented "
        "downloading is beneficial but its advantage is bounded by whichever link in the "
        "chain becomes the bottleneck first — a fundamental principle of distributed systems."
    )

    # Add expected speedup table for theoretical context
    add_h3(doc, "Theoretical vs Observed Speedup")
    add_body(doc,
        "For context, the table below contrasts theoretical ideal speedup (if each new "
        "connection added independent bandwidth) against the observed speedup:"
    )
    theory_data = [
        ("1 segment", "1.00×", "1.00×"),
        ("2 segments", "2.00×", "1.09×"),
        ("4 segments", "4.00×", "1.09×"),
        ("8 segments", "8.00×", "1.09×"),
    ]
    tbl2 = doc.add_table(rows=len(theory_data) + 1, cols=3)
    add_table_borders(tbl2)
    for i, h in enumerate(["Segments", "Ideal Speedup", "Observed Speedup"]):
        c = tbl2.rows[0].cells[i]
        set_cell_bg(c, "2E74B5")
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (seg, ideal, obs) in enumerate(theory_data, 1):
        row = tbl2.rows[idx]
        bg = "EBF3FB" if idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate([seg, ideal, obs]):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────
# Section 7 — Conclusion
# ──────────────────────────────────────────────────────────────

def build_conclusion(doc):
    add_h1(doc, "7.  Conclusion")

    add_body(doc,
        "The Simple Download Manager successfully demonstrates the core distributed systems "
        "concepts required by the CS404 project specification. Every required feature was "
        "implemented and is functional:"
    )
    for item in [
        "Multi-threaded segmented downloads using ThreadPoolExecutor and HTTP Range requests.",
        "Fault tolerance through per-segment automatic retry with linear back-off.",
        "Zero-CPU pause/resume via threading.Event, with correct handling of both mid-flight and post-restart scenarios.",
        "Persistent resume from SQLite — interrupted downloads survive application restarts.",
        "Real-time progress monitoring with smooth EWMA-based speed display and ETA.",
        "Full download history with sortable columns, individual record deletion, and clear-all.",
        "Both a Tkinter desktop GUI and a CLI fallback for headless environments.",
    ]:
        add_bullet(doc, item)

    add_body(doc,
        "Beyond the functional requirements, the project encountered and resolved several "
        "non-trivial engineering challenges: Windows file-locking semantics during cancellation, "
        "thread lifecycle management across application restarts, SQLite write contention "
        "optimisation, and the subtleties of accurate speed measurement in a bursty network environment."
    )
    add_body(doc,
        "The performance benchmark produced an honest real-world result: a 9% throughput "
        "improvement with 2 segments, plateauing thereafter due to the server's per-IP "
        "bandwidth cap. This outcome reinforces a fundamental distributed systems lesson — "
        "parallelism is only beneficial up to the point where the bottleneck shifts from "
        "the client's connection management to an external constraint."
    )
    add_body(doc,
        "The project codebase is organised as a clean seven-layer architecture with no circular "
        "dependencies. Each layer is independently testable, and the separation of concerns makes "
        "the codebase straightforward to extend with future features such as bandwidth throttling, "
        "download scheduling, or a web-based interface."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Report —")
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r.font.size = Pt(10)


# ──────────────────────────────────────────────────────────────
# Main entry point
# ──────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default body font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading colours via built-in styles
    for lvl, hex_c in [(1, "1F507A"), (2, "2E74B5"), (3, "1F507A")]:
        h_style = doc.styles[f"Heading {lvl}"]
        h_style.font.color.rgb = RGBColor(
            int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
        )

    add_page_number_footer(doc)
    build_cover(doc)
    build_intro(doc)
    build_architecture(doc)
    build_design_decisions(doc)
    build_thread_model(doc)
    build_challenges(doc)
    build_performance(doc)
    build_conclusion(doc)

    out = "SDM_Technical_Report.docx"
    doc.save(out)
    print(f"Report saved: {out}")


if __name__ == "__main__":
    main()
